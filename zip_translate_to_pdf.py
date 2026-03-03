import argparse
import os
import re
import sys
import textwrap
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import List, Tuple

# Readers
from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

# PDF
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib import colors

# Translation (free-ish; uses unofficial Google endpoint internally)
try:
    from deep_translator import GoogleTranslator
    HAS_TRANSLATOR = True
except Exception:
    HAS_TRANSLATOR = False


@dataclass
class Section:
    title: str
    body: str


def normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def read_html(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "html.parser")
    # remove scripts/styles
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text("\n")


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    # basic table capture
    for t in doc.tables:
        parts.append("\n[Table]")
        for row in t.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts).strip()


def read_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts).strip()


def translate_text(text: str, target_lang: str) -> str:
    """
    target_lang examples (deep-translator):
      'en' 'fr' 'de' 'es' 'ja' 'zh-CN' etc.
    """
    if not text.strip():
        return ""
    if not HAS_TRANSLATOR:
        return text  # no translation available; return original
    translator = GoogleTranslator(source="auto", target=target_lang)

    # chunk to avoid request limits
    max_chars = 4000
    chunks = [text[i:i + max_chars] for i in range(0, len(text), max_chars)]
    out = []
    for c in chunks:
        out.append(translator.translate(c))
    return "\n".join(out)


def extract_zip(zip_path: Path, out_dir: Path) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(zip_path, "r") as z:
        z.extractall(out_dir)


def load_sections(root: Path, target_lang: str, do_translate: bool) -> List[Section]:
    sections: List[Section] = []
    supported = {".txt", ".md", ".log", ".csv", ".html", ".htm", ".docx", ".pdf"}

    for path in sorted(root.rglob("*")):
        if path.is_dir():
            continue
        ext = path.suffix.lower()
        if ext not in supported:
            continue

        try:
            if ext in {".txt", ".md", ".log", ".csv"}:
                text = read_txt(path)
            elif ext in {".html", ".htm"}:
                text = read_html(path)
            elif ext == ".docx":
                text = read_docx(path)
            elif ext == ".pdf":
                text = read_pdf_text(path)
            else:
                continue
        except Exception as e:
            sections.append(Section(title=str(path.relative_to(root)), body=f"[Could not read: {e}]"))
            continue

        text = normalize(text)
        if not text:
            body = "[No extractable text]"
        else:
            body = translate_text(text, target_lang) if do_translate else text

        sections.append(Section(title=str(path.relative_to(root)), body=body))

    if not sections:
        sections.append(Section(title="(No supported files found)", body=""))
    return sections


def build_pdf(sections: List[Section], out_pdf: Path) -> None:
    out_pdf.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(
        str(out_pdf),
        pagesize=LETTER,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title=out_pdf.name,
        author="zip_translate_to_pdf.py",
    )

    styles = getSampleStyleSheet()
    title_style = styles["Heading1"]
    title_style.textColor = colors.HexColor("#1f2937")
    body_style = styles["BodyText"]
    body_style.leading = 14

    story = []
    story.append(Paragraph("Consolidated Document", styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Files included: {len(sections)}", styles["BodyText"]))
    story.append(PageBreak())

    # Simple TOC-like list
    story.append(Paragraph("Contents", styles["Heading2"]))
    for i, s in enumerate(sections, start=1):
        story.append(Paragraph(f"{i}. {s.title}", styles["BodyText"]))
    story.append(PageBreak())

    for i, s in enumerate(sections, start=1):
        story.append(Paragraph(f"{i}. {s.title}", title_style))
        story.append(Spacer(1, 8))

        # keep paragraphs reasonably sized for ReportLab
        chunks = s.body.split("\n\n")
        for para in chunks:
            para = para.strip()
            if not para:
                continue
            # Soft wrap very long lines
            para = "\n".join(textwrap.fill(line, width=120) for line in para.splitlines())
            story.append(Paragraph(para.replace("\n", "<br/>"), body_style))
            story.append(Spacer(1, 10))

        story.append(PageBreak())

    doc.build(story)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("zipfile", type=str, help="Path to input .zip")
    ap.add_argument("output_pdf", type=str, help="Path to output .pdf")
    ap.add_argument("--target", type=str, default="en", help="Target language code (default: en)")
    ap.add_argument("--no-translate", action="store_true", help="Only consolidate; no language translation")
    args = ap.parse_args()

    zip_path = Path(args.zipfile).expanduser().resolve()
    if not zip_path.exists():
        print(f"ZIP not found: {zip_path}")
        sys.exit(2)

    work_dir = zip_path.parent / f"{zip_path.stem}__unzipped"
    extract_zip(zip_path, work_dir)

    do_translate = not args.no_translate
    if do_translate and not HAS_TRANSLATOR:
        print("NOTE: deep-translator not available; will consolidate without translation.", file=sys.stderr)
        do_translate = False

    sections = load_sections(work_dir, target_lang=args.target, do_translate=do_translate)
    build_pdf(sections, Path(args.output_pdf).expanduser().resolve())
    print(f"Done. Wrote: {Path(args.output_pdf).expanduser().resolve()}")


if __name__ == "__main__":
    main()